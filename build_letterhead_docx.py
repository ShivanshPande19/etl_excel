"""
Generate an editable Word (.docx) letterhead for Crownest Hospitality LLP.

The company logo, name and GSTIN sit in the page HEADER, and the contact
details sit in the page FOOTER, so they repeat automatically on every page
while the body stays fully editable for typing letters.

Run:  python3 build_letterhead_docx.py
Output: crownest-letterhead.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Brand colours ----
MAROON = RGBColor(0x6B, 0x1F, 0x2A)
GOLD = RGBColor(0xC8, 0xA0, 0x4A)
GREY = RGBColor(0x6B, 0x72, 0x80)

LOGO_PATH = "logo/crownest-crest.png"


def set_cell_border_bottom(cell, color="6B1F2A", size="18"):
    """Add a bottom border to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    tcPr.append(borders)


def set_cell_border_top(cell, color="6B1F2A", size="18"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), color)
    borders.append(top)
    tcPr.append(borders)


def add_run(paragraph, text, size=11, bold=False, color=None, font="Georgia",
            spacing=None, italic=False):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color
    if spacing is not None:
        rPr = run._element.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rPr.append(sp)
    return run


def build():
    doc = Document()

    # ---- Page margins ----
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    # ================= HEADER =================
    header = section.header
    header.is_linked_to_previous = False
    # clear default empty paragraph
    htbl = header.add_table(rows=1, cols=3, width=Inches(6.7))
    htbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    htbl.autofit = False
    widths = [Inches(1.1), Inches(3.9), Inches(1.7)]
    for i, w in enumerate(widths):
        htbl.columns[i].width = w

    logo_cell, name_cell, gstin_cell = htbl.rows[0].cells
    for c in (logo_cell, name_cell, gstin_cell):
        for i, w in enumerate(widths):
            pass
        c.width = widths[[logo_cell, name_cell, gstin_cell].index(c)]

    # logo
    lp = logo_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        lp.add_run().add_picture(LOGO_PATH, width=Inches(0.95))
    except Exception:
        add_run(lp, "CROWNEST", size=12, bold=True, color=MAROON)

    # company name
    np1 = name_cell.paragraphs[0]
    np1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(np1, "CROWNEST ", size=22, bold=True, color=MAROON, font="Georgia")
    add_run(np1, "HOSPITALITY", size=22, bold=True, color=GOLD, font="Georgia")
    np2 = name_cell.add_paragraph()
    np2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(np2, "LIMITED LIABILITY PARTNERSHIP", size=8.5, bold=False,
            color=GOLD, font="Arial", spacing=40)
    np1.paragraph_format.space_after = Pt(2)

    # gstin
    gp1 = gstin_cell.paragraphs[0]
    gp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(gp1, "GSTIN", size=9, bold=True, color=MAROON, font="Arial")
    gp2 = gstin_cell.add_paragraph()
    gp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(gp2, "09AAVFC2225N1ZX", size=9, color=GREY, font="Arial")
    gp3 = gstin_cell.add_paragraph()
    gp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(gp3, "Greater Noida, U.P.", size=9, color=GREY, font="Arial")
    for p in (gp1, gp2):
        p.paragraph_format.space_after = Pt(1)

    # maroon bottom border under header (apply to all cells in the row)
    for c in htbl.rows[0].cells:
        set_cell_border_bottom(c, color="6B1F2A", size="18")

    # spacing line after header
    sp = header.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    add_run(sp, "", size=2)

    # ================= FOOTER =================
    footer = section.footer
    footer.is_linked_to_previous = False
    ftbl = footer.add_table(rows=1, cols=3, width=Inches(6.7))
    ftbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ftbl.autofit = False
    fwidths = [Inches(2.4), Inches(2.5), Inches(1.8)]
    office_cell, contact_cell, partners_cell = ftbl.rows[0].cells
    for i, w in enumerate(fwidths):
        ftbl.columns[i].width = w
    office_cell.width, contact_cell.width, partners_cell.width = fwidths

    def footer_block(cell, label, lines):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        add_run(p, label.upper(), size=7.5, bold=True, color=GOLD,
                font="Arial", spacing=30)
        for ln in lines:
            lp = cell.add_paragraph()
            lp.paragraph_format.space_after = Pt(0)
            add_run(lp, ln, size=8.5, color=GREY, font="Arial")

    footer_block(office_cell, "Registered Office", [
        "Site 5, G-107, Surajpur Kasna Road,",
        "Industrial Area, Greater Noida,",
        "Gautam Buddha Nagar, U.P. \u2013 201312",
    ])
    footer_block(contact_cell, "Contact", [
        "Tel: +91 99997 42752",
        "info@crownesthospitality.com",
        "abhishek@crownesthospitality.com",
        "puneet@crownesthospitality.com",
    ])
    footer_block(partners_cell, "Designated Partners", [
        "Abhishek Jindal",
        "Puneet Singh Anand",
    ])

    # maroon top border above footer
    for c in ftbl.rows[0].cells:
        set_cell_border_top(c, color="6B1F2A", size="18")

    # ================= BODY =================
    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_before = Pt(6)
    add_run(date_p, "Date: _______________", size=11, color=RGBColor(0, 0, 0),
            font="Georgia")

    doc.add_paragraph()  # blank line
    tip = doc.add_paragraph()
    add_run(tip, "Type your letter here\u2026", size=11, italic=True,
            color=RGBColor(0xB4, 0xB4, 0xB4), font="Georgia")

    # a few blank lines for writing room
    for _ in range(3):
        doc.add_paragraph()

    doc.save("crownest-letterhead.docx")
    print("Saved crownest-letterhead.docx")


if __name__ == "__main__":
    build()
