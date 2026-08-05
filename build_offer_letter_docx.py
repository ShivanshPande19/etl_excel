#!/usr/bin/env python3
"""
Build the Offer Letter content as a plain Word (.docx) document.

Just the letter content (with the original placeholders) - no letterhead.

Run:  python3 build_offer_letter_docx.py
Out:  Offer_Letter_Shivansh_Pande.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "Offer_Letter_Shivansh_Pande.docx"

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def para(text="", bold=False, align=None, space_after=6, size=11):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    return p


# --- Date ---
para("Date: [Insert Date]")
para()

# --- To block ---
para("To:")
para("Shivansh Pande", bold=True)
para("[Insert Your Address/City]")
para()

# --- Subject ---
para("Subject: Offer of Employment", bold=True)
para()

# --- Body ---
para("Dear Shivansh,")
para(
    "We are pleased to offer you the position of [Designation] at Crownest "
    "Hospitality LLP (a subsidiary of Azimuth Business on Wheels Private "
    "Limited)."
)
para("Please find the initial terms of your offer below:")

# Bulleted terms
doc.add_paragraph("Designation: [Designation]", style="List Bullet")
doc.add_paragraph("Duration: [Duration]", style="List Bullet")
doc.add_paragraph(
    "Compensation: Your salary will be [Salary Amount], subject to standard "
    "statutory deductions.",
    style="List Bullet",
)

para("Key Terms:", bold=True)

# Numbered key terms
doc.add_paragraph(
    'You will be expected to carry out your duties diligently, maintaining our '
    'commitment to "Hospitality with Integrity".',
    style="List Number",
)
doc.add_paragraph(
    "During your tenure, you will be bound by the standard rules, regulations, "
    "and confidentiality policies of Crownest Hospitality LLP and its parent "
    "company.",
    style="List Number",
)

para(
    "Please review this offer letter. If you accept the terms, kindly sign "
    "below and return a copy to us by [Insert Date]."
)
para("We look forward to welcoming you to our team.")
para()

# --- Signature block ---
para("Sincerely,")
para()
para("[Authorized Signatory Name]")
para("[Signatory Designation]")
para("Crownest Hospitality LLP", bold=True)

doc.save(OUTPUT)
print(f"Wrote {OUTPUT}")
