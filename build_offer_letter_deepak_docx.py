#!/usr/bin/env python3
"""
Build the Offer Letter content for Deepak Kumar as a plain Word (.docx).

Just the letter content - no letterhead (mirrors build_offer_letter_docx.py).

Run:  python3 build_offer_letter_deepak_docx.py
Out:  Offer_Letter_Deepak_Kumar.docx
"""

from docx import Document
from docx.shared import Pt

OUTPUT = "Offer_Letter_Deepak_Kumar.docx"

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def para(text="", bold=False, align=None, space_after=6, size=11):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    return p


# --- Date / Ref ---
para("Date: [Insert Date]")
para("Ref. No.: [Insert Ref. No.]")
para()

# --- To block ---
para("To:")
para("Deepak Kumar", bold=True)
para("Vill-Parthala Radha Kunj Colony Shiv Mandir,", space_after=0)
para("Post-Gautam Buddh Nagar,", space_after=0)
para("Distt.-Gautam Buddh Nagar,", space_after=0)
para("Uttar Pradesh - 201301", space_after=0)
para("Mobile No.: 9958821146")
para()

# --- Subject ---
para("Subject: Offer of Employment", bold=True)
para()

# --- Body ---
para("Dear Deepak,")
para(
    "We are pleased to offer you the position of House Keeper at Crownest "
    "Hospitality LLP (a subsidiary of Azimuth Business on Wheels Private "
    "Limited)."
)
para("Please find the initial terms of your offer below:")

# Bulleted terms
doc.add_paragraph("Designation: House Keeper", style="List Bullet")
doc.add_paragraph(
    "Compensation: Your salary will be Rs. 16,500/- per month, subject to "
    "standard statutory deductions.",
    style="List Bullet",
)

para("Key Terms:", bold=True)

# Numbered key terms
doc.add_paragraph(
    'You will be expected to carry out your duties diligently, maintaining our '
    'commitment to "Hospitality with Integrity".',
    style="List Number",
)
doc.add_paragraph(
    "During your tenure, you will be bound by the standard rules, regulations, "
    "and confidentiality policies of Crownest Hospitality LLP and its parent "
    "company.",
    style="List Number",
)

para(
    "Please review this offer letter. If you accept the terms, kindly sign "
    "below and return a copy to us by [Insert Date]."
)
para("We look forward to welcoming you to our team.")
para()

# --- Signature block ---
para("Sincerely,")
para()
para("[Authorized Signatory Name]")
para("[Signatory Designation]")
para("Crownest Hospitality LLP", bold=True)

doc.save(OUTPUT)
print(f"Wrote {OUTPUT}")
